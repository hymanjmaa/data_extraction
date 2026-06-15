"""
文档处理工具模块
支持PDF、图片、Word、Excel等文档解析
"""

import os
import pdfplumber
from PIL import Image
from typing import Dict, Any, List
import json
import base64
import io
import requests

class DocumentParser:
    """文档解析器基类"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, Any]:
        """解析PDF文档"""
        try:
            text_content = []
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

                pdf_text = '\n'.join(text_content)
                text_threshold = int(os.getenv('PDF_OCR_TEXT_THRESHOLD', '200'))
                need_ocr = (not pdf_text.strip()) or (len(pdf_text.strip()) < text_threshold)
                if need_ocr:
                    if os.getenv('PDF_OCR_ENABLED', '0') == '1':
                        ocr_provider = (os.getenv('PDF_OCR_PROVIDER', 'auto') or 'auto').strip().lower()
                        if ocr_provider in ['', 'auto']:
                            if os.getenv('BAILIAN_API_KEY'):
                                ocr_provider = 'bailian'
                            elif os.getenv('OPENROUTER_API_KEY'):
                                ocr_provider = 'openrouter'
                            else:
                                ocr_provider = ''

                        api_key = ''
                        base_url = ''
                        model = ''
                        extra_headers = {}

                        if ocr_provider == 'bailian':
                            api_key = os.getenv('BAILIAN_API_KEY', '')
                            base_url = os.getenv('BAILIAN_BASE_URL', 'https://dashscope.aliyuncs.com/compatible-mode/v1')
                            model = os.getenv('BAILIAN_OCR_MODEL', 'qwen-vl-plus')
                        elif ocr_provider == 'openrouter':
                            api_key = os.getenv('OPENROUTER_API_KEY', '')
                            base_url = os.getenv('OPENROUTER_BASE_URL', 'https://openrouter.ai/api/v1')
                            model = os.getenv('OPENROUTER_OCR_MODEL', '')
                        else:
                            api_key = ''

                        if not api_key:
                            return {
                                'success': False,
                                'error': 'PDF无可用文本（疑似扫描件），OCR未配置可用的API密钥（可设置PDF_OCR_PROVIDER=openrouter并配置OPENROUTER_OCR_MODEL，或使用百炼OCR）',
                                'type': 'pdf'
                            }
                        if ocr_provider == 'openrouter' and not model:
                            return {
                                'success': False,
                                'error': 'PDF无可用文本（疑似扫描件），请配置OPENROUTER_OCR_MODEL为支持图片输入的模型',
                                'type': 'pdf'
                            }
                        max_pages = int(os.getenv('PDF_OCR_MAX_PAGES', '5'))
                        resolution = int(os.getenv('PDF_OCR_RESOLUTION', '180'))

                        headers = {
                            'Authorization': f'Bearer {api_key}',
                            'Content-Type': 'application/json'
                        }

                        headers.update(extra_headers)
                        ocr_texts = []
                        for page in pdf.pages[:max_pages]:
                            img = page.to_image(resolution=resolution).original
                            buf = io.BytesIO()
                            img.save(buf, format='PNG')
                            b64 = base64.b64encode(buf.getvalue()).decode()


                            data = {
                                'model': model,
                                'messages': [
                                    {
                                        'role': 'user',
                                        'content': [
                                            {'type': 'text', 'text': '请识别图片中的文字，只输出识别出的文本，不要解释。'},
                                            {'type': 'image_url', 'image_url': {'url': f'data:image/png;base64,{b64}'}}
                                        ]
                                    }
                                ],
                                'temperature': 0.0,
                                'max_tokens': 2000
                            }

                            resp = requests.post(
                                f'{base_url}/chat/completions',
                                headers=headers,
                                json=data,
                                timeout=int(os.getenv('LLM_HTTP_TIMEOUT', '90'))
                            )
                            resp.raise_for_status()
                            result = resp.json()
                            content = result['choices'][0]['message']['content']
                            if content:
                                ocr_texts.append(content)

                        if not ocr_texts:
                            return {
                                'success': False,
                                'error': 'PDF无可用文本（疑似扫描件），OCR未识别到文字',
                                'type': 'pdf'
                            }

                        return {
                            'success': True,
                            'text': '\n'.join(ocr_texts),
                            'pages': len(ocr_texts),
                            'type': 'pdf_ocr'
                        }

                    return {
                        'success': False,
                        'error': 'PDF可提取文本过少（疑似扫描件），请开启OCR或提供可复制文本的PDF',
                        'type': 'pdf'
                    }

            return {
                'success': True,
                'text': '\n'.join(text_content),
                'pages': len(text_content),
                'type': 'pdf'
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'type': 'pdf'
            }
    
    @staticmethod
    def parse_image(file_path: str) -> Dict[str, Any]:
        """解析图片文档（OCR预处理）"""
        try:
            image = Image.open(file_path)
            
            return {
                'success': True,
                'image_data': file_path,
                'width': image.width,
                'height': image.height,
                'format': image.format,
                'type': 'image'
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'type': 'image'
            }
    
    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """解析Word文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs]
            tables = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                'success': True,
                'text': '\n'.join(paragraphs),
                'tables': tables,
                'type': 'docx'
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'type': 'docx'
            }
    
    @staticmethod
    def parse_xlsx(file_path: str) -> Dict[str, Any]:
        """解析Excel文档"""
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(file_path)
            sheets_data = {}
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                rows = []
                
                for row in sheet.iter_rows(values_only=True):
                    rows.append([str(cell) if cell is not None else '' for cell in row])
                
                sheets_data[sheet_name] = rows
            
            return {
                'success': True,
                'sheets': sheets_data,
                'sheet_names': wb.sheetnames,
                'type': 'xlsx'
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'type': 'xlsx'
            }
    
    @staticmethod
    def auto_parse(file_path: str) -> Dict[str, Any]:
        """自动识别并解析文档"""
        ext = os.path.splitext(file_path)[1].lower()
        
        parsers = {
            '.pdf': DocumentParser.parse_pdf,
            '.png': DocumentParser.parse_image,
            '.jpg': DocumentParser.parse_image,
            '.jpeg': DocumentParser.parse_image,
            '.docx': DocumentParser.parse_docx,
            '.xlsx': DocumentParser.parse_xlsx,
        }
        
        parser = parsers.get(ext)
        
        if parser:
            return parser(file_path)
        else:
            return {
                'success': False,
                'error': f'不支持的文件类型: {ext}',
                'type': ext
            }
